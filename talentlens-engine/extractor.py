import os
import json
import re

import ollama
import pymupdf as fitz  # PyMuPDF — `import fitz` directly is deprecated, `import pymupdf as fitz` is the current form
import pymupdf4llm
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pdf2image import convert_from_path
import pytesseract

from common import new_id, now_iso
from experience import compute_total_years

# Setting the data source to read directly from a local Kaggle directory path
DATA_DIR = os.path.expanduser("~/kaggle/input/talentlens-batch/")

MIN_TEXT_DENSITY_CHARS_PER_PAGE = 40  # below this, assume scanned/image PDF, fall back to OCR


# --- Stage 1: deterministic extraction (no LLM) ---

def _iter_block_items(doc: Document):
    """
    Yields paragraphs and tables in true document order. python-docx's .paragraphs and
    .tables are separate flat lists with no ordering between them, so a resume with e.g.
    a skills table in the middle of prose would come out with the table's content moved
    to the end — this walks the underlying XML body directly to preserve real order.
    """
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)


def _docx_to_markdown(file_path: str) -> str:
    """Structure-aware DOCX -> markdown: headings become #, list items become -, tables
    become markdown tables — all exact text, no summarization or rephrasing."""
    doc = Document(file_path)
    lines = []

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style_name = (block.style.name or "").lower()
            if "heading" in style_name or "title" in style_name:
                level = next((int(ch) for ch in style_name if ch.isdigit()), 1)
                lines.append(f"{'#' * min(max(level, 1), 6)} {text}")
            elif "list" in style_name:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        elif isinstance(block, Table):
            rows = block.rows
            if not rows:
                continue
            header = [c.text.strip() for c in rows[0].cells]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join("---" for _ in header) + " |")
            for row in rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


def extract_text(file_path: str) -> tuple[str, list[str], str]:
    """
    Extracts exact text from PDF, DOCX, or TXT files as markdown (PDF, DOCX) or plain
    text (TXT) — deterministic, no LLM involved at this stage.
    Returns (raw_text, extraction_warnings, extraction_method).
    PDFs with little/no extractable text (scanned/image PDFs) fall back to OCR.
    """
    ext = os.path.splitext(file_path)[1].lower()
    warnings = []

    if ext == '.pdf':
        with fitz.open(file_path) as doc:
            num_pages = doc.page_count

        raw_text = pymupdf4llm.to_markdown(file_path)
        avg_density = len(raw_text) / max(num_pages, 1)

        if avg_density >= MIN_TEXT_DENSITY_CHARS_PER_PAGE:
            return raw_text, warnings, "pdf_markdown"

        # Fall back to OCR — low text density strongly suggests a scanned/image PDF
        warnings.append("Low text density in PDF text layer — used OCR fallback.")
        images = convert_from_path(file_path)
        ocr_parts = [pytesseract.image_to_string(img) for img in images]
        ocr_text = "\n".join(ocr_parts).strip()
        if not ocr_text:
            warnings.append("OCR fallback also produced no text — file may be blank, corrupt, or unreadable.")
        return ocr_text, warnings, "pdf_ocr"

    elif ext == '.docx':
        raw_text = _docx_to_markdown(file_path)
        if not raw_text or not raw_text.strip():
            warnings.append("No extractable text found in DOCX — file may be empty or malformed.")
        return raw_text, warnings, "docx_markdown"

    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text, warnings, "txt"

    else:
        raise ValueError(f"Unsupported file type: {ext}")


# --- Stage 2: LLM structuring only (exact text -> JSON, no arithmetic) ---

def extract_structured_evidence(resume_text):
    """Sends already-exact resume markdown to Qwen to extract structured JSON. The model's
    only job is re-formatting text that's already in front of it — it does not compute
    durations or a total, and is not asked to."""

    system_prompt = (
        "You are an expert AI Resume Extractor. Your task is to extract information from the provided resume "
        "and output it EXACTLY matching the following JSON schema. Do not include any conversational text outside the JSON.\n\n"
        "IMPORTANT RULES:\n"
        "- Extract EVERY item explicitly listed under each resume section (every skill, every experience entry, "
        "every project, every certification, every education entry) — do not stop early or summarize a subset.\n"
        "- Only place an item in \"projects\" if it appears under a section literally titled Projects (or similar). "
        "Do NOT create a project entry from an Experience/Internship bullet point, and do not duplicate the same "
        "content across two fields.\n"
        "- For each experience entry, copy start_date_raw and end_date_raw EXACTLY as written in the resume text "
        "(e.g. \"Oct 2023\", \"Mar 2022\", \"2019\"). Do NOT calculate a duration, a number of years, or a total — "
        "that is computed separately, outside your output.\n"
        "- If the resume says the role is ongoing — using ANY wording such as \"Present\", \"Current\", \"Till "
        "date\", \"Ongoing\", or similar — copy that exact word as end_date_raw. This is a literal copy of text "
        "that IS on the page, not a guess. Only use an empty string for end_date_raw in the rare case where the "
        "resume shows a start date but truly no end date and no ongoing-indicator word at all.\n"
        "- For each experience entry, also extract technologies_used: the specific tools/technologies/platforms "
        "listed for THAT role specifically — e.g. a line like \"Tools and Techniques Used – Python, Docker, "
        "Kubernetes\" appearing under that job. Keep this separate from the top-level skills list below.\n"
        "- For each project entry, extract technologies_used the same way — tools/technologies mentioned for "
        "that specific project. Keep this separate from both the top-level skills list and from any experience "
        "entry's technologies_used.\n"
        "- The top-level \"skills\" field is ONLY for items listed under a resume section literally titled "
        "Skills / Technical Skills / Skills and Technologies (or similar) — not items pulled from job or project "
        "descriptions, even if they'd fit there too.\n"
        "- If a field has no information in the resume, use an empty array or empty string — never invent content.\n\n"
        "SCHEMA:\n"
        "{\n"
        '  "candidate_name": "string (extracted from resume)",\n'
        '  "skills": ["string", "string"],\n'
        '  "experience": [\n'
        '    {\n'
        '      "title": "string",\n'
        '      "company": "string",\n'
        '      "start_date_raw": "string — copied exactly as written, e.g. \'Oct 2023\'",\n'
        '      "end_date_raw": "string — copied exactly as written, e.g. \'Present\' or \'Mar 2022\'",\n'
        '      "domain": "string (inferred domain/industry)",\n'
        '      "description": "string",\n'
        '      "technologies_used": ["string", "string"]\n'
        '    }\n'
        '  ],\n'
        '  "education": [\n'
        '    { "degree_level": "string", "field": "string", "institution": "string" }\n'
        '  ],\n'
        '  "certifications": ["string", "string"],\n'
        '  "projects": [\n'
        '    { "title": "string", "description": "string", "technologies_used": ["string", "string"] }\n'
        '  ]\n'
        "}"
    )

    response = ollama.chat(
        model='qwen2.5:3b-instruct',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': f"Resume Text:\n{resume_text}"}
        ],
        format='json',
        # No think=False here — qwen2.5 doesn't have qwen3's thinking-token mode,
        # so there's nothing to disable.
        options={
            'temperature': 0.0,
            'num_predict': 8192
        }
    )

    raw_output = response['message']['content']

    try:
        return json.loads(raw_output)
    except json.JSONDecodeError:
        print("\n[WARNING] Model failed to return a complete JSON object. Returning raw text.")
        return {"error": "Incomplete JSON", "raw_output": raw_output}


# --- Stage 2b: completeness check + targeted backfill ---
#
# The full structured-extraction call above can silently drop an entry even on a short,
# fully-visible input (confirmed: a 4.3k-char resume, well under the 8192 num_predict
# budget, still lost its 3rd/last job and the last line of its Skills section). That
# rules out truncation — it's the small instruct model paying less attention to the tail
# of a single long multi-section extraction task than the head, not a budget problem.
# num_predict/prompt wording can't reliably fix that, so instead of trusting the one big
# call, we independently verify it against a much simpler, low-complexity list-only call
# (much less prone to the same drop-under-load failure), and re-extract anything missing
# with a narrow, single-target follow-up call.

def _list_companies_present(resume_text: str) -> list[str]:
    """Cheap completeness-check pass: ask ONLY for company names under Work Experience,
    nothing else. A short list-only task is much less likely to drop an item than the
    full multi-field structured extraction, so this serves as an independent check on
    len(structured["experience"])."""
    system_prompt = (
        "List every company name that appears under the Work Experience / Employment "
        "History section of this resume, one per line, in the exact order they appear. "
        "Company names ONLY — no job titles, no dates, no other text. Include every "
        "entry, even short-tenure, older, or internship roles. Do not skip any."
    )
    response = ollama.chat(
        model='qwen2.5:3b-instruct',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': f"Resume Text:\n{resume_text}"}
        ],
        options={'temperature': 0.0, 'num_predict': 512}
    )
    lines = response['message']['content'].splitlines()
    companies = [ln.strip(" -*•\t").strip() for ln in lines]
    return [c for c in companies if c]


_NON_COMPANY_LABELS = {
    "phone", "email", "address", "contact", "contact info", "contact information",
    "mobile", "location", "linkedin", "github", "portfolio", "website", "skype",
    "personal details", "languages", "hobbies", "interests", "certifications",
    "skills", "education", "summary", "objective", "references", "projects",
    "date of birth", "nationality", "marital status", "technical skills",
    "professional skills", "work experience", "experience",
}


def _looks_like_non_company_label(name: str) -> bool:
    """True if `name` is a resume section/contact-info label rather than a real company
    name (e.g. 'Phone', 'Email', 'Address'). Confirmed necessary in production: a
    two-column resume layout can flatten a CONTACT sidebar (Phone/Email/Address labels)
    into text near Work Experience, and the company-listing LLM pass can mistake those
    labels for employer names — which then causes _reextract_single_job to fabricate a
    full experience entry (title, description, technologies_used) for an employer that
    doesn't exist. Exact match only, after normalization — a real company literally
    named e.g. 'Phone Inc' would normalize to 'phone inc', not 'phone', so this can't
    exclude it. This is a first line of defense; _reextract_single_job's company-match
    check below is the second, for cases this denylist doesn't happen to cover."""
    return _norm_company(name) in _NON_COMPANY_LABELS


def _companies_match(a: str, b: str) -> bool:
    """Fuzzy company-name match (substring either direction, on normalized text) — same
    rule already used elsewhere to decide 'is this company already extracted'."""
    na, nb = _norm_company(a), _norm_company(b)
    return bool(na) and bool(nb) and (na in nb or nb in na)


def _reextract_single_job(resume_text: str, company_name: str) -> dict | None:
    """Targeted single-entry extraction for one company the main pass dropped. Much
    narrower scope than the full-resume call, so it's far less likely to drop the one
    thing it's being asked for. Returns the entry dict, or None if the model still
    couldn't produce it (left for manual review rather than guessed) OR if the model
    returned an entry for a DIFFERENT company than the one requested — confirmed in
    production: asked to recover 'ML-Based Fuel Delivery Automation & Optimization'
    (actually a project subheading inside the Foster Technologies India block, not a
    company — see _looks_like_non_company_label for the contact-label variant of the
    same root problem), the model didn't fail; it returned a fabricated entry under
    company='Foster Technologies India' instead — a real, ALREADY-extracted company —
    which silently created a duplicate experience entry with no warning that anything
    was wrong, since entry.get('company') was truthy. Validating the returned company
    actually matches what was asked closes that gap."""
    schema_snippet = (
        '{"title": "string", "company": "string", '
        '"start_date_raw": "string — copied exactly as written", '
        '"end_date_raw": "string — copied exactly as written, or the ongoing-indicator '
        'word if present", "domain": "string", "description": "string", '
        '"technologies_used": ["string", "string"]}'
    )
    system_prompt = (
        f"From the resume text below, extract ONLY the work experience entry for the "
        f"company \"{company_name}\". Output a single JSON object matching this schema "
        f"exactly, with no other text:\n{schema_snippet}\n\n"
        "Copy dates and text exactly as written. Do not summarize, invent, or omit fields. "
        f"If \"{company_name}\" does NOT actually appear as an employer in the Work "
        "Experience section (e.g. it's a section label, a project name, or contact "
        "info rather than a company), output a JSON object with an empty string for "
        "every field instead of guessing or substituting a different company."
    )
    response = ollama.chat(
        model='qwen2.5:3b-instruct',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': f"Resume Text:\n{resume_text}"}
        ],
        format='json',
        options={'temperature': 0.0, 'num_predict': 1024}
    )
    try:
        entry = json.loads(response['message']['content'])
    except json.JSONDecodeError:
        return None

    if not isinstance(entry, dict):
        return None
    returned_company = entry.get("company")
    if not returned_company:
        return None
    if not _companies_match(returned_company, company_name):
        # Model returned SOMETHING, but not for the company we asked about — treat
        # exactly like a failed recovery rather than silently accepting a mismatched
        # (often duplicate) entry.
        return None

    return entry


def _extract_technologies_for_job(resume_text: str, company_name: str) -> list[str]:
    """Dedicated, narrow LLM pass for JUST the tools/technologies list of one specific
    job. Same rationale as _extract_skills_section: _reextract_single_job asks for
    title, company, both dates, domain, description, AND technologies_used in one JSON
    object — technologies_used is the LAST field, and a narrow single-purpose call is
    less prone to that tail-of-a-structured-task drop than the multi-field call is.

    Confirmed in production this closes the TAIL-drop gap but is not by itself fully
    reliable against MID-list drops (see _extract_technologies_deterministic) — this is
    why _backfill_technologies unions this with a deterministic regex pass rather than
    trusting this alone."""
    system_prompt = (
        f"The resume text below contains a Work Experience entry for the company "
        f"\"{company_name}\" (often listed under a line like 'Tools and Techniques Used' "
        f"or similar, specific to that role). Extract EVERY individual tool, technology, "
        f"or platform listed for that specific role — do not skip any, especially ones at "
        "the end of the list. Output one item per line, no bullets, no numbering, no other "
        "text. If that role lists no tools/technologies, output nothing."
    )
    response = ollama.chat(
        model='qwen2.5:3b-instruct',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': f"Resume Text:\n{resume_text}"}
        ],
        options={'temperature': 0.0, 'num_predict': 512}
    )
    lines = response['message']['content'].splitlines()
    items = [ln.strip(" -*•\t").strip() for ln in lines]
    return [i for i in items if i]


_TECH_LABEL_RE = re.compile(
    r'(?:tools?\s*(?:and|&)?\s*techniques?\s*used|tech(?:nical)?\s*stack|'
    r'technologies\s*(?:used|and\s*techniques))\s*[-–—:]\s*(.+)',
    re.IGNORECASE | re.DOTALL  # DOTALL: the tools list commonly wraps across a line break
    # in the source PDF (e.g. "...Power\nBI") — without DOTALL, "." stops at the first
    # newline and silently truncates the captured list mid-item instead of un-wrapping it.
)


def _extract_technologies_deterministic(resume_text: str, company_name: str, all_companies: list) -> list[str]:
    """Best-effort, no-LLM extraction of a 'Tools and Techniques Used –' style line for one
    company's section. Purely additive alongside _extract_technologies_for_job: if the
    resume doesn't use this heading style, or the section boundary can't be located,
    this returns an empty list and nothing else is affected.

    Why this exists: confirmed in production on Kamna Kashyap's Zucol Group entry, the
    narrow single-purpose LLM pass (_extract_technologies_for_job) still silently dropped
    two items — Computer Vision, YOLO — from the MIDDLE of a 9-item comma list, not the
    tail. That's a different failure mode than the tail-drop the narrow pass was designed
    to fix, and prompt/temperature tuning can't guarantee it won't recur. Where the resume
    explicitly labels this list (which pymupdf4llm preserves as exact text — no OCR/LLM
    involved), a regex extraction of that exact line has no such failure mode at all.

    Section boundary: from company_name's first occurrence up to whichever OTHER company
    name appears next in the text (or end of text) — a rough but effective proxy for
    "this job's section", since resumes list one company's tools line under that company.
    """
    company_pos = resume_text.lower().find(company_name.strip().lower())
    if company_pos == -1:
        return []

    section_end = len(resume_text)
    for other in all_companies:
        if _norm_company(other) == _norm_company(company_name):
            continue
        pos = resume_text.lower().find(other.strip().lower(), company_pos + len(company_name))
        if pos != -1 and pos < section_end:
            section_end = pos

    section_text = resume_text[company_pos:section_end]
    match = _TECH_LABEL_RE.search(section_text)
    if not match:
        return []

    tail = match.group(1)
    # Cut off at the first BLANK line or markdown heading-ish marker, so a wrapped/loose
    # match doesn't pull in unrelated text following the tools line. A single newline
    # (no blank line) is treated as a wrapped continuation, not a delimiter — the source
    # PDF wraps mid-item (e.g. "Power\nBI"), so splitting on every newline would cut a
    # single item in half.
    tail = re.split(r'\n\s*\n|\n#|\n\*\*[A-Z]', tail)[0]
    tail = tail.replace('\n', ' ')
    items = tail.split(',')
    return [i.strip(" -*•\t.") for i in items if i.strip(" -*•\t.")]


def _backfill_technologies(resume_text: str, entry: dict, company_name: str, all_companies: list = None) -> str | None:
    """Runs BOTH the dedicated LLM technologies-only pass and the deterministic regex pass,
    and merges anything either one found (that the recovery pass missed) into
    entry["technologies_used"] in place. Case-insensitive dedup; preserves the recovery
    pass's existing items/casing/order and appends anything genuinely new after — same
    merge shape as _fill_missing_skills. Union of two independent signals rather than
    either alone, since production confirmed the LLM pass by itself isn't fully reliable
    (see _extract_technologies_deterministic docstring). Returns a warning string
    describing what was recovered, or None if nothing new was found."""
    recovered_items = _extract_technologies_for_job(resume_text, company_name)
    recovered_items += _extract_technologies_deterministic(resume_text, company_name, all_companies or [])

    existing = entry.get("technologies_used", []) or []
    existing_norm = {t.strip().lower() for t in existing if isinstance(t, str) and t.strip()}

    newly_added = []
    for item in recovered_items:
        norm = item.strip().lower()
        if norm and norm not in existing_norm:
            existing.append(item)
            existing_norm.add(norm)
            newly_added.append(item)

    entry["technologies_used"] = existing

    if not newly_added:
        return None
    return (
        f"{len(newly_added)} technology item(s) for '{company_name}' were missing from "
        f"the job-recovery pass and backfilled via a dedicated technologies-only "
        f"re-extraction pass: {', '.join(newly_added)}. Verify manually."
    )


def _fill_missing_experience(resume_text: str, structured: dict) -> list[str]:
    """Runs the completeness check and backfills any missing experience entries in
    place (mutates structured["experience"]). Returns warning strings for anything
    recovered or anything that couldn't be recovered."""
    warnings = []
    expected_companies = _list_companies_present(resume_text)
    extracted = structured.get("experience", []) or []
    extracted_companies_norm = [(_norm_company(e.get("company", ""))) for e in extracted]

    for company in expected_companies:
        norm = _norm_company(company)
        if not norm:
            continue
        if _looks_like_non_company_label(company):
            # Silently skip, no warning: this is the expected/common case (a contact
            # sidebar label or section header flattened into the Work Experience text
            # by the PDF layout), not something worth flagging for manual review every
            # time. See _looks_like_non_company_label — confirmed necessary in
            # production (Phone/Email/Address mistaken for employers).
            continue
        already_present = any(norm in ec or ec in norm for ec in extracted_companies_norm if ec)
        if already_present:
            continue

        recovered = _reextract_single_job(resume_text, company)
        if recovered:
            tech_warning = _backfill_technologies(resume_text, recovered, company, expected_companies)
            structured.setdefault("experience", []).append(recovered)
            warnings.append(
                f"'{company}' was missing from the initial extraction and was recovered "
                f"via a targeted follow-up pass. Please verify its fields manually."
            )
            if tech_warning:
                warnings.append(tech_warning)
        else:
            warnings.append(
                f"'{company}' was flagged as a possible employer not in the initial "
                f"extraction, but the follow-up recovery pass could not confirm it as a "
                f"real employer entry (or found one under a different, already-known "
                f"company — a possible sign '{company}' is actually a project name or "
                f"section label rather than an employer). Verify manually; do not "
                f"assume total_years_experience is affected without checking."
            )

    return warnings


def _norm_company(name: str) -> str:
    return (name or "").strip().lower()


# --- Stage 2c: Skills-section completeness backfill ---
#
# Same failure mode as the dropped Zucol Group job — confirmed on a resume where the
# raw text (verified directly, not inferred) contains the Skills section's "Algorithms"
# sub-line intact, but it never appeared anywhere in structured["skills"], on a resume
# well under the token budget. Skills doesn't need the detect-then-recover two-step the
# Experience fix uses, though: it's a flat list with no ambiguous "which one is missing"
# matching problem, so a single dedicated section-scoped extraction pass — cheap, and
# narrow enough to reliably return every sub-category — can just always run, then union +
# dedup its output into structured["skills"].

def _extract_skills_section(resume_text: str) -> list[str]:
    """Dedicated, section-scoped pass for JUST the Skills/Technical Skills section. A
    narrower list-extraction task is much less likely to drop tail items than the full
    multi-section extraction — isolating this section keeps the call small enough that it
    reliably returns every sub-category (Languages, Frameworks, Platforms, Algorithms,
    Tools, etc.), all flattened into one list, same shape as the main "skills" field."""
    system_prompt = (
        "The resume text below contains a Skills / Technical Skills / Skills and "
        "Technologies section. That section is commonly broken into multiple "
        "sub-categories with their own labels (e.g. Languages, Frameworks, Platforms, "
        "Algorithms, Tools, Databases). Extract EVERY individual skill, technology, "
        "framework, or algorithm item listed anywhere in that section, across ALL "
        "sub-categories including the LAST one — do not stop early and do not skip any "
        "sub-category, especially the final one. Output one item per line, no category "
        "headers, no bullets, no numbering, no other text."
    )
    response = ollama.chat(
        model='qwen2.5:3b-instruct',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': f"Resume Text:\n{resume_text}"}
        ],
        options={'temperature': 0.0, 'num_predict': 1024}
    )
    lines = response['message']['content'].splitlines()
    items = [ln.strip(" -*•\t").strip() for ln in lines]
    return [i for i in items if i]


def _fill_missing_skills(resume_text: str, structured: dict) -> list[str]:
    """Runs the dedicated skills-section pass and merges anything the main extraction
    missed into structured["skills"] in place. Case-insensitive dedup; preserves the main
    pass's existing items/casing/order and appends anything genuinely new after. Returns
    warning strings describing what was recovered, for extraction_warnings."""
    warnings = []
    recovered_items = _extract_skills_section(resume_text)

    existing = structured.get("skills", []) or []
    existing_norm = {s.strip().lower() for s in existing if isinstance(s, str) and s.strip()}

    newly_added = []
    for item in recovered_items:
        norm = item.strip().lower()
        if norm and norm not in existing_norm:
            existing.append(item)
            existing_norm.add(norm)
            newly_added.append(item)

    structured["skills"] = existing

    if newly_added:
        warnings.append(
            f"{len(newly_added)} skill item(s) were missing from the initial extraction "
            f"and recovered via a dedicated skills-section re-extraction pass: "
            f"{', '.join(newly_added)}. Verify manually."
        )

    return warnings


def _aggregate_skills(structured: dict) -> list:
    """Merges the resume's own Skills-section list with every technologies_used entry
    from experience roles and projects, deduped case-insensitively, preserving first-seen
    casing/order. Kept as a separate field (skills_all_sources) rather than overwriting
    "skills" so "skills" still faithfully reflects just the resume's own Skills section —
    useful for display/explainability even after scoring uses the broader set."""
    seen = set()
    aggregated = []

    def _add(item):
        if not item or not isinstance(item, str):
            return
        key = item.strip().lower()
        if key and key not in seen:
            seen.add(key)
            aggregated.append(item)

    for skill in structured.get("skills", []) or []:
        _add(skill)
    for entry in structured.get("experience", []) or []:
        for tech in entry.get("technologies_used", []) or []:
            _add(tech)
    for proj in structured.get("projects", []) or []:
        for tech in proj.get("technologies_used", []) or []:
            _add(tech)

    return aggregated


def ingest_resume(file_path: str) -> dict:
    """
    Full resume ingestion:
      Stage 1 — extract exact text (pymupdf4llm / docx-to-markdown, OCR fallback)
      Stage 2 — structure via LLM (text/dates only, no arithmetic)
      Stage 3 — compute total_years_experience deterministically from the extracted
                date ranges (experience.py), via range union
    then attach candidate_id/document_id/file_name/extraction metadata.
    """
    file_name = os.path.basename(file_path)
    raw_text, warnings, method = extract_text(file_path)
    structured = extract_structured_evidence(raw_text)

    # Completeness check + backfill — must run BEFORE compute_total_years, since a
    # recovered job's dates directly affect the union-total calculation (this is exactly
    # what caused Kamna Kashyap's total to read 4.5 yrs instead of ~5.7 yrs).
    completeness_warnings = _fill_missing_experience(raw_text, structured)
    warnings = warnings + completeness_warnings

    # Same completeness problem, applied to the Skills section (confirmed: the "Algorithms"
    # sub-line was silently dropped even though it's present in the raw text). Must run
    # BEFORE _aggregate_skills so skills_all_sources picks up any recovered items too.
    skills_completeness_warnings = _fill_missing_skills(raw_text, structured)
    warnings = warnings + skills_completeness_warnings

    total_years, experience_warnings = compute_total_years(structured.get("experience", []))
    structured["total_years_experience"] = total_years
    structured["skills_all_sources"] = _aggregate_skills(structured)
    warnings = warnings + experience_warnings

    structured["candidate_id"] = new_id()
    structured["document_id"] = new_id()
    structured["file_name"] = file_name
    structured["extraction_method"] = method
    structured["extraction_warnings"] = warnings
    structured["uploaded_at"] = now_iso()
    return structured


if __name__ == "__main__":
    sample_resume = os.path.join(DATA_DIR, "Naukri_VibhutiSharma[5y_6m].pdf")

    if os.path.exists(sample_resume):
        print("Ingesting resume...")
        structured_data = ingest_resume(sample_resume)
        print(json.dumps(structured_data, indent=2))
    else:
        print(f"Please ensure your test resume is placed in: {DATA_DIR}")