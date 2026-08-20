import os
import re
import json
import logging

import ollama
import pymupdf as fitz  # PyMuPDF — `import fitz` directly is deprecated, `import pymupdf as fitz` is the current form
import pymupdf4llm
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pdf2image import convert_from_path
import pytesseract

from common import new_id, now_iso
from experience import compute_total_years

logger = logging.getLogger("talentlens.extractor")

# Setting the data source to read directly from a local Kaggle directory path
DATA_DIR = os.path.expanduser("~/kaggle/input/talentlens-batch/")

MIN_TEXT_DENSITY_CHARS_PER_PAGE = 40  # below this, assume scanned/image PDF, fall back to OCR


# --- Stage 1: deterministic extraction (no LLM) ---

def _iter_block_items(doc: Document):
    """
    Yields paragraphs and tables in true document order. python-docx's .paragraphs and
    .tables are separate flat lists with no ordering between them, so a resume with e.g.
    a skills table in the middle of prose would come out with the table's content moved
    to the end — this walks the underlying XML body directly to preserve real order.
    """
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)


def _docx_to_markdown(file_path: str) -> str:
    """Structure-aware DOCX -> markdown: headings become #, list items become -, tables
    become markdown tables — all exact text, no summarization or rephrasing."""
    doc = Document(file_path)
    lines = []

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style_name = (block.style.name or "").lower()
            if "heading" in style_name or "title" in style_name:
                level = next((int(ch) for ch in style_name if ch.isdigit()), 1)
                lines.append(f"{'#' * min(max(level, 1), 6)} {text}")
            elif "list" in style_name:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        elif isinstance(block, Table):
            rows = block.rows
            if not rows:
                continue
            header = [c.text.strip() for c in rows[0].cells]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join("---" for _ in header) + " |")
            for row in rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


# --- Stage 1 addition: column-aware page reordering for multi-column PDFs ---

_FULL_WIDTH_FRACTION = 0.6        # block wider than this fraction of page width = header/divider, not a column
_MIN_COLUMN_BLOCKS = 4            # fewer candidate blocks than this = not enough evidence to call it two-column
_MIN_COLUMN_GAP_FRACTION = 0.03   # required horizontal gap between the two x-clusters, as a fraction of page width
_MIN_COLUMN_HEIGHT_OVERLAP = 0.3  # left/right block y-ranges must overlap by at least this fraction of page height,
                                   # so a real running sidebar is required — not just one late block with a different x


def _cluster_columns_1d(xs: list) -> tuple | None:
    """Simple 1D two-means clustering on block left-x (x0) positions."""
    if len(xs) < _MIN_COLUMN_BLOCKS:
        return None

    lo, hi = min(xs), max(xs)
    if hi - lo < 1e-6:
        return None

    left_c, right_c = lo, hi
    for _ in range(20):
        left_pts = [x for x in xs if abs(x - left_c) <= abs(x - right_c)]
        right_pts = [x for x in xs if abs(x - left_c) > abs(x - right_c)]
        if not left_pts or not right_pts:
            return None  # collapsed to one cluster -- not two-column
        new_left = sum(left_pts) / len(left_pts)
        new_right = sum(right_pts) / len(right_pts)
        converged = abs(new_left - left_c) < 0.5 and abs(new_right - right_c) < 0.5
        left_c, right_c = new_left, new_right
        if converged:
            break

    return (left_c, right_c) if left_c < right_c else (right_c, left_c)


def _page_two_column_split_x(narrow_blocks: list, page_width: float, page_height: float):
    """Returns the x-coordinate to split narrow_blocks into left/right columns."""
    if len(narrow_blocks) < _MIN_COLUMN_BLOCKS:
        return None

    centers = _cluster_columns_1d([b[0] for b in narrow_blocks])
    if centers is None:
        return None
    left_c, right_c = centers

    if (right_c - left_c) < _MIN_COLUMN_GAP_FRACTION * page_width:
        return None

    split_x = (left_c + right_c) / 2
    left_blocks = [b for b in narrow_blocks if b[0] < split_x]
    right_blocks = [b for b in narrow_blocks if b[0] >= split_x]
    if len(left_blocks) < 2 or len(right_blocks) < 2:
        return None

    left_y_span = (min(b[1] for b in left_blocks), max(b[3] for b in left_blocks))
    right_y_span = (min(b[1] for b in right_blocks), max(b[3] for b in right_blocks))
    overlap = min(left_y_span[1], right_y_span[1]) - max(left_y_span[0], right_y_span[0])
    if overlap < _MIN_COLUMN_HEIGHT_OVERLAP * page_height:
        return None

    return split_x


def _reorder_page_columns_aware(page) -> str:
    """Returns column-corrected text for one fitz page."""
    page_width, page_height = page.rect.width, page.rect.height
    raw_blocks = [b for b in page.get_text("blocks") if b[4].strip() and b[6] == 0]

    full_width = [b for b in raw_blocks if (b[2] - b[0]) > _FULL_WIDTH_FRACTION * page_width]
    narrow = [b for b in raw_blocks if (b[2] - b[0]) <= _FULL_WIDTH_FRACTION * page_width]

    split_x = _page_two_column_split_x(narrow, page_width, page_height)
    if split_x is None:
        return page.get_text("text")

    left = sorted((b for b in narrow if b[0] < split_x), key=lambda b: b[1])
    right = sorted((b for b in narrow if b[0] >= split_x), key=lambda b: b[1])

    narrow_top = min(b[1] for b in narrow)
    header = sorted((b for b in full_width if b[1] < narrow_top), key=lambda b: b[1])
    trailer = sorted((b for b in full_width if b[1] >= narrow_top), key=lambda b: b[1])

    ordered = [b[4].strip() for b in header + left + right + trailer]
    return "\n\n".join(t for t in ordered if t)


def _extract_pdf_text_column_aware(file_path: str):
    """Whole-document column-aware extraction."""
    try:
        with fitz.open(file_path) as doc:
            any_two_column = False
            pages_out = []
            for page in doc:
                page_width, page_height = page.rect.width, page.rect.height
                raw_blocks = [b for b in page.get_text("blocks") if b[4].strip() and b[6] == 0]
                narrow = [b for b in raw_blocks if (b[2] - b[0]) <= _FULL_WIDTH_FRACTION * page_width]
                if _page_two_column_split_x(narrow, page_width, page_height) is not None:
                    any_two_column = True
                    pages_out.append(_reorder_page_columns_aware(page))
                else:
                    pages_out.append(page.get_text("text"))
    except Exception as e:
        logger.warning("Column-aware PDF extraction failed, falling back to pymupdf4llm: %s", e)
        return None

    if not any_two_column:
        return None

    return "\n\n".join(pages_out)


def extract_text(file_path: str) -> tuple[str, list[str], str]:
    """
    Extracts exact text from PDF, DOCX, or TXT files as markdown (PDF, DOCX) or plain
    text (TXT) — deterministic, no LLM involved at this stage.
    """
    ext = os.path.splitext(file_path)[1].lower()
    warnings = []

    if ext == '.pdf':
        with fitz.open(file_path) as doc:
            num_pages = doc.page_count

        column_aware_text = _extract_pdf_text_column_aware(file_path)
        if column_aware_text is not None:
            warnings.append(
                "Detected a multi-column page layout (e.g. a sidebar running "
                "alongside a main column) — used column-aware block reordering "
                "instead of pymupdf4llm's default reading order, to avoid "
                "interleaving sidebar content with main-column content."
            )
            raw_text = column_aware_text
        else:
            raw_text = pymupdf4llm.to_markdown(file_path)

        avg_density = len(raw_text) / max(num_pages, 1)

        if avg_density >= MIN_TEXT_DENSITY_CHARS_PER_PAGE:
            return raw_text, warnings, "pdf_markdown"

        warnings.append("Low text density in PDF text layer — used OCR fallback.")
        images = convert_from_path(file_path)
        ocr_parts = [pytesseract.image_to_string(img) for img in images]
        ocr_text = "\n".join(ocr_parts).strip()
        if not ocr_text:
            warnings.append("OCR fallback also produced no text — file may be blank, corrupt, or unreadable.")
        return ocr_text, warnings, "pdf_ocr"

    elif ext == '.docx':
        raw_text = _docx_to_markdown(file_path)
        if not raw_text or not raw_text.strip():
            warnings.append("No extractable text found in DOCX — file may be empty or malformed.")
        return raw_text, warnings, "docx_markdown"

    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text, warnings, "txt"

    else:
        raise ValueError(f"Unsupported file type: {ext}")


# --- Stage 2: LLM structuring only (exact text -> JSON, no arithmetic) ---

def _attempt_json_repair(raw_output: str):
    """Best-effort salvage of a malformed model response, tried only after
    json.loads() has already failed."""
    try:
        from json_repair import repair_json
    except ImportError:
        logger.warning(
            "json_repair package not installed — cannot salvage truncated model output. "
            "Run: pip install json-repair"
        )
        return None

    try:
        repaired = repair_json(raw_output, return_objects=True)
        if isinstance(repaired, dict) and repaired:
            return repaired
    except Exception as e:
        logger.warning("json_repair raised while salvaging output: %s", e)
    return None


def _dedupe_repeated_experience(structured: dict) -> int:
    """Collapses experience entries that are exact duplicates."""
    entries = structured.get("experience") or []
    if not entries:
        return 0

    def _key(e):
        return (
            (e.get("title") or "").strip().lower(),
            (e.get("company") or "").strip().lower(),
            (e.get("start_date_raw") or "").strip().lower(),
            (e.get("end_date_raw") or "").strip().lower(),
            re.sub(r"\s+", " ", (e.get("description") or "").strip().lower()),
        )

    seen, deduped, removed = set(), [], 0
    for e in entries:
        k = _key(e)
        if k in seen:
            removed += 1
            continue
        seen.add(k)
        deduped.append(e)

    structured["experience"] = deduped
    return removed


_STRUCTURED_EVIDENCE_SCHEMA = {
    "type": "object",
    "properties": {
        "candidate_name": {"type": "string"},
        "current_role_title_from_summary": {"type": "string"},
        "stated_years_experience_from_summary": {"type": "number"},
        "skills": {"type": "array", "items": {"type": "string"}, "maxItems": 80},
        "experience": {
            "type": "array",
            "maxItems": 15,
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "company": {"type": "string"},
                    "start_date_raw": {"type": "string"},
                    "end_date_raw": {"type": "string"},
                    "domain": {"type": "string"},
                    "technologies_used": {"type": "array", "items": {"type": "string"}, "maxItems": 40},
                },
                "required": ["title", "company", "start_date_raw", "end_date_raw"],
            },
        },
        "education": {
            "type": "array",
            "maxItems": 10,
            "items": {
                "type": "object",
                "properties": {
                    "degree_level": {"type": "string"},
                    "field": {"type": "string"},
                    "institution": {"type": "string"},
                },
            },
        },
        "certifications": {"type": "array", "items": {"type": "string"}, "maxItems": 25},
        "projects": {
            "type": "array",
            "maxItems": 20,
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "technologies_used": {"type": "array", "items": {"type": "string"}, "maxItems": 40},
                },
            },
        },
    },
    "required": ["candidate_name", "skills", "experience", "education", "certifications", "projects"],
}


def extract_structured_evidence(resume_text):
    """Sends already-exact resume markdown to Qwen to extract structured JSON."""

    system_prompt = (
        "You are an expert AI Resume Extractor. Your task is to extract information from the provided resume "
        "and output it EXACTLY matching the following JSON schema. Do not include any conversational text outside the JSON.\n\n"
        "IMPORTANT RULES:\n"
        "- Extract EVERY item explicitly listed under each resume section (every skill, every experience entry, "
        "every project, every certification, every education entry) — do not stop early or summarize a subset.\n"
        "- Only place an item in \"projects\" if it appears under a section literally titled Projects (or similar). "
        "Do NOT create a project entry from an Experience/Internship bullet point, and do not duplicate the same "
        "content across two fields.\n"
        "- For each experience entry, copy start_date_raw and end_date_raw EXACTLY as written in the resume text "
        "(e.g. \"Oct 2023\", \"Mar 2022\", \"2019\"). Do NOT calculate a duration, a number of years, or a total — "
        "that is computed separately, outside your output.\n"
        "- If the resume says the role is ongoing — using ANY wording such as \"Present\", \"Current\", \"Till "
        "date\", \"Ongoing\", or similar — copy that exact word as end_date_raw. This is a literal copy of text "
        "that IS on the page, not a guess. Only use an empty string for end_date_raw in the rare case where the "
        "resume shows a start date but truly no end date and no ongoing-indicator word at all.\n"
        "- For each experience entry, also extract technologies_used: the specific tools/technologies/platforms "
        "listed for THAT role specifically — e.g. a line like \"Tools and Techniques Used – Python, Docker, "
        "Kubernetes\" appearing under that job. Keep this separate from the top-level skills list below.\n"
        "- For each project entry, extract technologies_used the same way — tools/technologies mentioned for "
        "that specific project. Keep this separate from both the top-level skills list and from any experience "
        "entry's technologies_used.\n"
        "- The top-level \"skills\" field is ONLY for items listed under a resume section literally titled "
        "Skills / Technical Skills / Skills and Technologies (or similar) — not items pulled from job or project "
        "descriptions, even if they'd fit there too.\n"
        "- An \"experience\" entry corresponds ONLY to a distinct job/role block under a Work Experience or "
        "Employment History section — one entry per unique (title, company, date range) combination actually "
        "printed on the resume. Do NOT create a separate experience entry for each category or grouping listed "
        "under Technical Skills / Skills (e.g. \"Data Engineering & Pipelines\", \"Programming & Frameworks\", "
        "\"Model Deployment\" are skill categories, never job roles).\n"
        "- Never repeat the same experience entry more than once. If you notice yourself about to write an "
        "entry with the same title, company, and dates as one you already wrote, stop — do not write it again.\n"
        "- If a field has no information in the resume, use an empty array or empty string — never invent content.\n\n"
        "CURRENT ROLE TITLE FROM SUMMARY (current_role_title_from_summary field):\n"
        "- If the resume has a Summary / Profile / About / Objective section that explicitly states the "
        "candidate's current or most recent job title in prose (e.g. \"Data Scientist with 5 years of "
        "experience in...\", \"Results-driven Senior Data Engineer specializing in...\"), copy JUST that "
        "title phrase exactly as written (e.g. \"Data Scientist\", \"Senior Data Engineer\").\n"
        "- This is separate from — and in addition to — the title you already extract for each entry in "
        "\"experience\" below. It exists because a resume's stated current title in prose doesn't always "
        "exactly match, or isn't always clearly the first/most-recent entry in, the dated experience list "
        "(unconventional ordering, a title only mentioned in prose, etc.).\n"
        "- Only extract this if the summary/profile text ACTUALLY states a title. Do NOT infer or guess a "
        "title from the candidate's skills, most recent employer, or general seniority level. If there is "
        "no summary/profile section, or it doesn't state a title, current_role_title_from_summary MUST be "
        "an empty string — that is correct and expected, not something to fill in.\n\n"
        "STATED YEARS OF EXPERIENCE FROM SUMMARY (stated_years_experience_from_summary field):\n"
        "- If the resume has a Summary / Profile / About section that explicitly states the candidate's "
        "total years of experience (e.g., \"5+ years of experience\", \"Over 8 years in data engineering\"), "
        "extract the numeric value (e.g., 5, 8).\n"
        "- Do NOT calculate or guess this number by looking at the job history. Only extract it if it is "
        "explicitly written in the summary prose. If it is not stated, output 0.\n\n"
        "SCHEMA:\n"
        "{\n"
        '  "candidate_name": "string (extracted from resume)",\n'
        '  "current_role_title_from_summary": "string — ONLY if the resume\'s summary/profile prose states '
        'a current/most-recent title; empty string otherwise",\n'
        '  "stated_years_experience_from_summary": "number — exact number stated in summary prose, or 0 if not stated",\n'
        '  "skills": ["string", "string"],\n'
        '  "experience": [\n'
        '    {\n'
        '      "title": "string",\n'
        '      "company": "string",\n'
        '      "start_date_raw": "string — copied exactly as written, e.g. \'Oct 2023\'",\n'
        '      "end_date_raw": "string — copied exactly as written, e.g. \'Present\' or \'Mar 2022\'",\n'
        '      "domain": "string (inferred domain/industry)",\n'
        '      "description": "string",\n'
        '      "technologies_used": ["string", "string"]\n'
        '    }\n'
        '  ],\n'
        '  "education": [\n'
        '    { "degree_level": "string", "field": "string", "institution": "string" }\n'
        '  ],\n'
        '  "certifications": ["string", "string"],\n'
        '  "projects": [\n'
        '    { "title": "string", "description": "string", "technologies_used": ["string", "string"] }\n'
        '  ]\n'
        "}"
    )

    response = ollama.chat(
        model='qwen2.5:3b-instruct',
        messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': f"Resume Text:\n{resume_text}"}
        ],
        format=_STRUCTURED_EVIDENCE_SCHEMA,
        options={
            'temperature': 0.0,
            'num_ctx': 8192,
        }
    )

    raw_output = response['message']['content']

    try:
        structured = json.loads(raw_output)
        removed = _dedupe_repeated_experience(structured)
        if removed:
            logger.warning(
                "Model repeated %d experience entr%s verbatim — collapsed via dedupe.",
                removed, "y" if removed == 1 else "ies"
            )
        return structured
    except json.JSONDecodeError:
        logger.warning("Model returned incomplete JSON — attempting salvage before discarding.")
        repaired = _attempt_json_repair(raw_output)
        if repaired is not None:
            removed = _dedupe_repeated_experience(repaired)
            note = f" ({removed} duplicate experience entr{'y' if removed == 1 else 'ies'} collapsed)" if removed else ""
            logger.warning("Recovered partial data from truncated model output via repair%s.", note)
            repaired.setdefault("_extraction_note", f"Recovered from truncated/malformed model output{note}.")
            return repaired
        print("\n[WARNING] Model failed to return a complete JSON object, and it could not be salvaged. Returning raw text.")
        return {"error": "Incomplete JSON", "raw_output": raw_output}


def _aggregate_skills(structured: dict) -> list:
    """Merges the resume's own Skills-section list with every technologies_used entry
    from experience roles and projects, deduped case-insensitively, preserving first-seen
    casing/order."""
    seen = set()
    aggregated = []

    def _add(item):
        if not item or not isinstance(item, str):
            return
        key = item.strip().lower()
        if key and key not in seen:
            seen.add(key)
            aggregated.append(item)

    for skill in structured.get("skills", []) or []:
        _add(skill)
    for entry in structured.get("experience", []) or []:
        for tech in entry.get("technologies_used", []) or []:
            _add(tech)
    for proj in structured.get("projects", []) or []:
        for tech in proj.get("technologies_used", []) or []:
            _add(tech)

    # NEW: Extract skills from certifications formatted as "Label: skill, skill"
    for cert in structured.get("certifications", []) or []:
        if ":" in cert:
            _, content = cert.split(":", 1)
            for token in _split_dense_skill_line(content):
                _add(token)

    return aggregated


# --- Fabrication guards and deterministic recovery ---

def _normalize_for_guard(text: str) -> str:
    """Normalizes whitespace AND typographic quote/dash characters before comparison."""
    text = (text or "")
    text = text.replace("\u201c", '"').replace("\u201d", '"')  # curly double quotes
    text = text.replace("\u2018", "'").replace("\u2019", "'")  # curly single quotes
    text = text.replace("\u2013", "-").replace("\u2014", "-")  # en/em dash
    return re.sub(r"\s+", " ", text.strip().lower())


def _text_supports_item(item: str, raw_text_norm: str) -> bool:
    """True if `item` is actually supported by the raw resume text."""
    item_norm = _normalize_for_guard(item)
    if not item_norm:
        return False

    pattern = r"(?<![a-z0-9])" + re.escape(item_norm) + r"(?![a-z0-9])"
    if re.search(pattern, raw_text_norm):
        return True

    words = item_norm.split()
    if len(words) > 1:
        return all(
            re.search(r"(?<![a-z0-9])" + re.escape(w) + r"(?![a-z0-9])", raw_text_norm)
            for w in words
        )
    return False


def _guard_against_fabricated_skills(structured: dict, raw_text: str) -> list:
    """Drops any "skills" or "certifications" entry that doesn't actually appear
    anywhere in the candidate's raw resume text."""
    warnings = []
    raw_text_norm = _normalize_for_guard(raw_text)

    for field in ("skills", "certifications"):
        items = structured.get(field) or []
        if not items:
            continue

        kept, dropped = [], []
        for item in items:
            if _text_supports_item(item, raw_text_norm):
                kept.append(item)
            else:
                dropped.append(item)

        if dropped:
            warnings.append(
                f"{field}: discarded {len(dropped)} entr{'y' if len(dropped) == 1 else 'ies'} "
                f"not found anywhere in the raw resume text ({', '.join(dropped)}) — likely "
                f"model fabrication rather than something actually on the page. Verify manually."
            )
            structured[field] = kept

    return warnings


def _guard_against_fabricated_summary_title(structured: dict, raw_text: str) -> list:
    """Clears current_role_title_from_summary if it isn't actually supported by the
    raw resume text."""
    title = (structured.get("current_role_title_from_summary") or "").strip()
    if not title:
        return []

    raw_text_norm = _normalize_for_guard(raw_text)
    if _text_supports_item(title, raw_text_norm):
        return []

    structured["current_role_title_from_summary"] = ""
    return [
        f"current_role_title_from_summary: discarded {title!r} — not found anywhere in the raw "
        f"resume text — likely model fabrication rather than something actually stated in a "
        f"summary/profile section. Verify manually."
    ]


_LABELED_SKILL_LINE_RE = re.compile(r"^\s*([A-Za-z][A-Za-z0-9 &/\-]{1,40}):\s*(.+)$")
_SKILL_LABEL_HINT_RE = re.compile(
    r"skill|technolog|tool|stack|language|framework|pipeline|engineering|database", re.IGNORECASE
)
_MID_PREPOSITION_RE = re.compile(r"\b(for|with|in|of|using|on)\b", re.IGNORECASE)
_MAX_RECOVERED_TOKEN_WORDS = 3


def _split_dense_skill_line(content: str) -> list:
    """Splits a colon-labeled skill line's content into candidate atomic tokens."""
    candidates = []
    for part in re.split(r",|&", content):
        part = re.sub(r"^(?:and|or)\s+", "", part.strip(" ."), flags=re.IGNORECASE).strip(" .")
        if not part or _MID_PREPOSITION_RE.search(part) or len(part.split()) > _MAX_RECOVERED_TOKEN_WORDS:
            continue
        candidates.append(part)
    return candidates


def _recover_labeled_skill_lines(raw_text: str, structured: dict) -> list:
    """Returns the skill tokens recovered from raw_text that Stage 2 dropped."""
    existing_norm = {_normalize_for_guard(s) for s in (structured.get("skills") or [])}
    for entry in structured.get("experience", []) or []:
        existing_norm.update(_normalize_for_guard(t) for t in entry.get("technologies_used", []) or [])
    for proj in structured.get("projects", []) or []:
        existing_norm.update(_normalize_for_guard(t) for t in proj.get("technologies_used", []) or [])

    recovered = []
    for line in (raw_text or "").splitlines():
        m = _LABELED_SKILL_LINE_RE.match(line.strip())
        if not m or not _SKILL_LABEL_HINT_RE.search(m.group(1)):
            continue
        for token in _split_dense_skill_line(m.group(2)):
            token_norm = _normalize_for_guard(token)
            if not token_norm or token_norm in existing_norm:
                continue
            existing_norm.add(token_norm)
            recovered.append(token)

    if recovered:
        structured["skills"] = (structured.get("skills") or []) + recovered
    return recovered


_COMPANY_DATE_LINE_RE = re.compile(
    r"^\s*([A-Z][A-Za-z0-9&.,'’\- ]{2,60}?)\s*\n\s*"
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4})"
    r"\s*[-–—]+\s*"
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|present|current|till date|ongoing)",
    re.IGNORECASE | re.MULTILINE,
)


def _check_experience_completeness(raw_text: str, structured: dict) -> list:
    """Returns warnings listing company names that look like a job header in
    raw_text but have no matching entry in structured['experience']."""
    text = raw_text or ""
    candidates = set()
    for m in _COMPANY_DATE_LINE_RE.finditer(text):
        name = re.sub(r"\s+", " ", m.group(1).strip())
        if len(name.split()) <= 6:
            candidates.add(name)

    extracted_companies_norm = {
        _normalize_for_guard(e.get("company") or "") for e in structured.get("experience", []) or []
    }

    missing = sorted(
        c for c in candidates
        if c and not any(_normalize_for_guard(c) in ec or ec in _normalize_for_guard(c)
                          for ec in extracted_companies_norm if ec)
    )

    if not missing:
        return []
    return [
        f"experience: found a company/date-range header in the raw resume text with no "
        f"matching entry in the extracted experience list ({', '.join(missing)}) — the "
        f"model may have dropped a job. Verify manually."
    ]


def ingest_resume(file_path: str) -> dict:
    """
    Full resume ingestion:
      Stage 1 — extract exact text (pymupdf4llm / docx-to-markdown, OCR fallback)
      Stage 2 — structure via LLM (text/dates only, no arithmetic)
      Stage 3 — compute total_years_experience deterministically from the extracted
                date ranges (experience.py), via range union
    then attach candidate_id/document_id/file_name/extraction metadata.
    """
    file_name = os.path.basename(file_path)
    raw_text, warnings, method = extract_text(file_path)

    _debug_dump_path = os.path.join("/tmp", f"raw_text_debug__{file_name}.txt")
    with open(_debug_dump_path, "w", encoding="utf-8") as _f:
        _f.write(f"[extraction_method={method}]\n[char_count={len(raw_text)}]\n\n{raw_text}")
    print(f"[DEBUG] Stage 1 raw_text ({len(raw_text)} chars, method={method}) written to {_debug_dump_path}")

    structured = extract_structured_evidence(raw_text)

    recovered_skills = _recover_labeled_skill_lines(raw_text, structured)
    if recovered_skills:
        warnings.append(
            f"skills: recovered {len(recovered_skills)} token(s) present in the raw resume "
            f"text but dropped by Stage 2 extraction ({', '.join(recovered_skills)})."
        )

    warnings = warnings + _guard_against_fabricated_skills(structured, raw_text)
    warnings = warnings + _guard_against_fabricated_summary_title(structured, raw_text)

    total_years, experience_warnings = compute_total_years(structured.get("experience", []))
    
    # NEW: Check if the summary explicitly stated a higher number of years
    stated_years = structured.get("stated_years_experience_from_summary", 0)
    
    if isinstance(stated_years, (int, float)) and stated_years > total_years:
        structured["total_years_experience"] = float(stated_years)
        warnings.append(f"experience: Used stated experience from summary ({stated_years}y) because it exceeded calculated experience ({total_years}y).")
    else:
        structured["total_years_experience"] = total_years
        
    structured["skills_all_sources"] = _aggregate_skills(structured)
    warnings = warnings + experience_warnings
    warnings = warnings + _check_experience_completeness(raw_text, structured)

    for w in warnings:
        logger.warning("[%s] %s", file_name, w)

    structured["candidate_id"] = new_id()
    structured["document_id"] = new_id()
    structured["file_name"] = file_name
    structured["extraction_method"] = method
    structured["extraction_warnings"] = warnings
    structured["uploaded_at"] = now_iso()
    return structured


if __name__ == "__main__":
    sample_resume = os.path.join(DATA_DIR, "Naukri_Ulka[5y_0m].pdf")

    if os.path.exists(sample_resume):
        print("Ingesting resume...")
        structured_data = ingest_resume(sample_resume)
        print(json.dumps(structured_data, indent=2))
    else:
        print(f"Please ensure your test resume is placed in: {DATA_DIR}")